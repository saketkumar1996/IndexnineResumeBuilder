from fastapi import APIRouter, HTTPException, UploadFile, File
from fastapi.responses import Response
from pydantic import BaseModel, ValidationError
from typing import List, Dict, Any, Optional
from models.resume_models import ResumeModel
from core.resume_normalizer import normalize_resume_input
import os
import json
import re
from datetime import datetime, timezone
from pathlib import Path

router = APIRouter()
AI_OUTPUT_LOG_DIR = Path(os.getenv("AI_OUTPUT_LOG_DIR", Path(__file__).resolve().parents[1] / "ai_output_logs"))
PROFESSIONAL_EXPERIENCE_BULLET_LIMIT = 3
PROJECT_EXPERIENCE_BULLET_LIMIT = 2


def _validated_resume_from_any_shape(resume_data: Dict[str, Any]) -> ResumeModel:
    return ResumeModel(**normalize_resume_input(resume_data))


def _filename_from_resume(resume_data: Dict[str, Any], extension: str) -> str:
    header = resume_data.get("header") or {}
    name = header.get("fullName") or header.get("name") or "resume"
    safe_name = re.sub(r"[^A-Za-z0-9_-]+", "_", name).strip("_") or "resume"
    return f"{safe_name}.{extension}"


def _add_docx_bullets(doc, items: List[str], limit: Optional[int] = None):
    for item in items[:limit] if limit else items:
        if str(item or "").strip():
            doc.add_paragraph(str(item).strip(), style="List Bullet")


def _generate_docx_bytes(resume_data: Dict[str, Any]) -> bytes:
    from docx import Document
    from io import BytesIO

    doc = Document()
    header = resume_data.get("header") or {}
    expertise = resume_data.get("expertise") or {}
    skills = resume_data.get("skills") or {}

    doc.add_heading(header.get("fullName") or header.get("name") or "Resume", 0)
    title = header.get("designation") or header.get("title")
    contact = " | ".join(
        value
        for value in [
            title,
            header.get("email"),
            header.get("phone"),
            header.get("location"),
            header.get("linkedin"),
            header.get("github"),
            header.get("portfolio"),
        ]
        if value
    )
    if contact:
        doc.add_paragraph(contact)

    if expertise.get("summary"):
        doc.add_heading("Professional Summary", level=1)
        doc.add_paragraph(expertise["summary"])
    _add_docx_bullets(doc, expertise.get("bulletPoints") or [])

    if skills.get("skills"):
        doc.add_heading("Skills", level=1)
        doc.add_paragraph(skills["skills"])

    experiences = resume_data.get("experiences") or resume_data.get("experience") or []
    if experiences:
        doc.add_heading("Professional Experience", level=1)
        for exp in experiences:
            company = exp.get("company", "")
            role = exp.get("title") or exp.get("position") or ""
            dates = " - ".join(value for value in [exp.get("startDate") or exp.get("start_date"), exp.get("endDate") or exp.get("end_date")] if value)
            doc.add_paragraph(f"{role} | {company} | {dates}".strip(" |"))
            _add_docx_bullets(doc, exp.get("responsibilities") or [], PROFESSIONAL_EXPERIENCE_BULLET_LIMIT)

    if resume_data.get("projects"):
        doc.add_heading("Selected Projects", level=1)
        for project in resume_data.get("projects") or []:
            if not project.get("name") and not project.get("description"):
                continue
            doc.add_paragraph(project.get("name") or "Project", style="List Bullet")
            if project.get("technologies"):
                doc.add_paragraph(f"Technologies: {project['technologies']}")
            if project.get("description"):
                doc.add_paragraph(project["description"])
            _add_docx_bullets(doc, project.get("responsibilities") or [], PROJECT_EXPERIENCE_BULLET_LIMIT)

    if resume_data.get("education"):
        doc.add_heading("Education", level=1)
        for edu in resume_data.get("education") or []:
            if edu.get("institution") or edu.get("degree"):
                dates = " - ".join(value for value in [edu.get("startYear"), edu.get("endYear")] if value)
                doc.add_paragraph(f"{edu.get('degree', '')} | {edu.get('institution', '')} | {dates}".strip(" |"))

    if resume_data.get("awards"):
        doc.add_heading("Awards & Certifications", level=1)
        for award in resume_data.get("awards") or []:
            if award.get("title"):
                doc.add_paragraph(
                    " | ".join(value for value in [award.get("title"), award.get("organization"), award.get("year")] if value),
                    style="List Bullet",
                )

    out = BytesIO()
    doc.save(out)
    return out.getvalue()


class ValidationResponse(BaseModel):
    """Response model for validation endpoint"""
    valid: bool
    errors: Optional[List[Dict[str, Any]]] = None
    data: Optional[Dict[str, Any]] = None


class PreviewResponse(BaseModel):
    """Response model for preview generation endpoint"""
    html: str
    valid: bool
    errors: Optional[List[Dict[str, Any]]] = None


@router.post("/validate", response_model=ValidationResponse)
async def validate_resume(resume_data: Dict[str, Any]) -> ValidationResponse:
    """
    Validate resume data against spec
    
    This endpoint validates the complete resume data structure using Pydantic models
    and returns structured validation results with field-specific errors.
    
    Requirements: 5.2, 5.3, 5.4
    """
    try:
        validated_resume = _validated_resume_from_any_shape(resume_data)
        
        return ValidationResponse(
            valid=True,
            data=validated_resume.model_dump(),
            errors=None
        )
        
    except ValidationError as e:
        # Format Pydantic validation errors for frontend consumption
        formatted_errors = []
        
        for error in e.errors():
            # Extract field path and error details
            field_path = ".".join(str(loc) for loc in error["loc"])
            error_msg = error["msg"]
            error_type = error["type"]
            
            formatted_error = {
                "field": field_path,
                "message": error_msg,
                "type": error_type,
                "input": error.get("input", None)
            }
            
            # Add spec reference based on error type
            if "word" in error_msg.lower():
                formatted_error["spec_reference"] = "Requirements 2.3"
            elif "date" in error_msg.lower() or "MMM YYYY" in error_msg:
                formatted_error["spec_reference"] = "Requirements 2.2"
            elif "responsibilities" in error_msg.lower():
                formatted_error["spec_reference"] = "Requirements 2.4"
            elif "emoji" in error_msg.lower() or "icon" in error_msg.lower():
                formatted_error["spec_reference"] = "Requirements 2.5"
            elif "comma" in error_msg.lower():
                formatted_error["spec_reference"] = "Requirements 6.2"
            
            formatted_errors.append(formatted_error)
        
        return ValidationResponse(
            valid=False,
            errors=formatted_errors,
            data=None
        )
    
    except Exception as e:
        # Handle unexpected errors
        raise HTTPException(
            status_code=500,
            detail=f"Internal server error during validation: {str(e)}"
        )


@router.post("/preview", response_model=PreviewResponse)
async def generate_preview(resume_data: Dict[str, Any]) -> PreviewResponse:
    """
    Generate HTML preview for validated data
    
    This endpoint first validates the resume data, then generates an HTML preview
    using Jinja2 templates that mirror the DOCX layout.
    
    Requirements: 3.2, 3.5
    """
    try:
        # First validate the data
        validated_resume = _validated_resume_from_any_shape(resume_data)
        
        # Import renderer here to avoid circular imports
        from rendering.resume_renderer import ResumeRenderer
        
        # Generate HTML preview
        renderer = ResumeRenderer(template_path="templates")
        html_content = renderer.render_html(validated_resume)
        
        return PreviewResponse(
            html=html_content,
            valid=True,
            errors=None
        )
        
    except ValidationError as e:
        # Generate partial preview with available data
        try:
            from rendering.resume_renderer import ResumeRenderer
            renderer = ResumeRenderer(template_path="templates")
            
            # Create a partial HTML preview showing what we have
            partial_html = renderer.render_partial_html(resume_data)
            
            # Format validation errors
            formatted_errors = []
            for error in e.errors():
                field_path = ".".join(str(loc) for loc in error["loc"])
                formatted_errors.append({
                    "field": field_path,
                    "message": error["msg"],
                    "type": error["type"]
                })
            
            return PreviewResponse(
                html=partial_html,
                valid=False,
                errors=formatted_errors
            )
            
        except Exception:
            # Fallback to error message
            formatted_errors = []
            for error in e.errors():
                field_path = ".".join(str(loc) for loc in error["loc"])
                formatted_errors.append({
                    "field": field_path,
                    "message": error["msg"],
                    "type": error["type"]
                })
            
            return PreviewResponse(
                html="<div class='error'>Validation errors prevent preview generation. Please fix the errors and try again.</div>",
                valid=False,
                errors=formatted_errors
            )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error generating preview: {str(e)}"
        )


@router.post("/export")
async def export_pdf(resume_data: Dict[str, Any]):
    """
    Generate PDF file for validated data
    
    This endpoint validates the resume data and generates a PDF file using
    WeasyPrint library with HTML template rendering.
    
    Requirements: 4.1, 4.5
    """
    try:
        # First validate the data
        validated_resume = _validated_resume_from_any_shape(resume_data)
        
        # Import renderer and FastAPI response classes
        from rendering.resume_renderer import ResumeRenderer
        
        # Generate PDF document
        renderer = ResumeRenderer(template_path="templates")
        pdf_bytes = renderer.generate_pdf(validated_resume)
        
        # Return PDF response
        return Response(
            content=pdf_bytes,
            media_type="application/pdf",
            headers={"Content-Disposition": "attachment; filename=resume.pdf"}
        )
        
    except ValidationError as e:
        # Return validation error - export should only proceed if validation passes
        raise HTTPException(
            status_code=422,
            detail={
                "message": "Validation failed - export blocked",
                "errors": [
                    {
                        "field": ".".join(str(loc) for loc in error["loc"]),
                        "message": error["msg"]
                    }
                    for error in e.errors()
                ]
            }
        )
    
    except Exception as e:
        raise HTTPException(
            status_code=500,
            detail=f"Error generating PDF export: {str(e)}"
        )


@router.post("/export/docx")
async def export_docx(resume_data: Dict[str, Any]):
    """Generate an editable DOCX from the current frontend resume shape."""
    try:
        docx_bytes = _generate_docx_bytes(resume_data)
        filename = _filename_from_resume(resume_data, "docx")
        return Response(
            content=docx_bytes,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename={filename}"},
        )
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error generating DOCX export: {str(e)}")


# Resume extraction prompt for AI
RESUME_SCHEMA_PROMPT = """You are a resume parser. Extract structured resume data from the uploaded resume document (PDF or DOCX).

Return ONLY valid JSON (no markdown, no code fences, no explanation). Use this exact structure:

{
  "header": {
    "fullName": "string",
    "designation": "string",
    "email": "string or empty",
    "phone": "string or empty",
    "location": "string or empty",
    "linkedin": "string or empty",
    "github": "string or empty",
    "portfolio": "string or empty"
  },
  "expertise": {
    "summary": "string, 80-120 words professional summary",
    "bulletPoints": ["string", "string"]
  },
  "skills": { "skills": "comma-separated string" },
  "experiences": [
    { "company": "", "title": "", "location": "", "startDate": "MMM YYYY", "endDate": "MMM YYYY or Present", "responsibilities": ["up to 3 resume bullets if present"] }
  ],
  "projects": [
    {
      "name": "",
      "client": "",
      "description": "",
      "technologies": "",
      "developmentTools": "",
      "teamSize": "",
      "responsibilities": ["up to 2 resume bullets if present"],
      "link": ""
    }
  ],
  "education": [
    { "institution": "", "degree": "", "location": "", "startYear": "YYYY", "endYear": "YYYY", "gpa": "", "honors": "" }
  ],
  "awards": [
    { "title": "", "year": "YYYY", "organization": "" }
  ]
}

Rules:
- Use empty string for missing scalar fields and [] for missing responsibility arrays.
- Dates: "Jan 2020", "Present". Years: "2015", "2019".
- Extract up to 3 real bullets for each Professional Experience entry if bullets are present. Do not create bullets when the resume has none.
- Extract all project descriptions and up to 2 real project responsibility bullets for each Selected Project if bullets are present.
- If a project has bullets but no explicit description paragraph, use the first project bullet as description and put the remaining bullets in responsibilities.
- Do not put project descriptions in responsibilities only.
- Extract everything you can find; omit arrays only if none found."""


def _extract_json_from_response(content: str) -> Dict[str, Any]:
    """Pull raw JSON out of AI response (handles markdown code blocks)."""
    content = content.strip()
    # Strip markdown code block if present
    m = re.search(r"```(?:json)?\s*([\s\S]*?)\s*```", content)
    if m:
        content = m.group(1).strip()
    return json.loads(content)


def _safe_log_filename(filename: str) -> str:
    safe_name = re.sub(r"[^A-Za-z0-9_.-]+", "_", filename).strip("._") or "resume"
    return safe_name[:80]


def _write_ai_output_log(
    *,
    upload_filename: str,
    raw_ai_response: str,
    parsed_data: Optional[Dict[str, Any]] = None,
    returned_data: Optional[Dict[str, Any]] = None,
    error: Optional[str] = None,
) -> Optional[Path]:
    """Write upload AI output for local debugging. The folder is gitignored."""
    try:
        AI_OUTPUT_LOG_DIR.mkdir(parents=True, exist_ok=True)
        timestamp = datetime.now(timezone.utc).strftime("%Y%m%dT%H%M%SZ")
        log_path = AI_OUTPUT_LOG_DIR / f"{timestamp}_{_safe_log_filename(upload_filename)}.json"
        log_path.write_text(
            json.dumps(
                {
                    "timestamp": timestamp,
                    "uploadFilename": upload_filename,
                    "rawAiResponse": raw_ai_response,
                    "parsedData": parsed_data,
                    "returnedData": returned_data,
                    "error": error,
                },
                indent=2,
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )
        return log_path
    except Exception:
        return None


def _clean_text(value: Any) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def _first_text(source: Dict[str, Any], *keys: str) -> str:
    for key in keys:
        value = source.get(key)
        if isinstance(value, list):
            value = next((_clean_text(item) for item in value if _clean_text(item)), "")
        text = _clean_text(value)
        if text:
            return text
    return ""


def _text_list(value: Any) -> List[str]:
    if isinstance(value, list):
        items = value
    elif isinstance(value, str) and ("\n" in value or "•" in value):
        items = re.split(r"[\n•]+", value)
    elif value:
        items = [value]
    else:
        items = []

    cleaned: List[str] = []
    for item in items:
        if isinstance(item, dict):
            item = _first_text(item, "text", "description", "summary", "responsibility", "title")
        text = re.sub(r"^[\s\-•*]+", "", _clean_text(item))
        if text:
            cleaned.append(text)
    return cleaned


def _first_text_list(source: Dict[str, Any], *keys: str) -> List[str]:
    for key in keys:
        values = _text_list(source.get(key))
        if values:
            return values
    return []


def _item_list(value: Any) -> List[Any]:
    if isinstance(value, list):
        return value
    if isinstance(value, dict):
        return [value]
    return []


def _normalize_uploaded_project(project: Any) -> Dict[str, Any]:
    source = project if isinstance(project, dict) else {}
    responsibilities = _first_text_list(
        source,
        "responsibilities",
        "responsibility",
        "bullets",
        "bulletPoints",
        "highlights",
        "achievements",
        "contributions",
        "tasks",
    )
    description = _first_text(source, "description", "projectDescription", "summary", "overview", "details")
    if not description and responsibilities:
        description = responsibilities[0]
        responsibilities = responsibilities[1:]
    responsibilities = responsibilities[:PROJECT_EXPERIENCE_BULLET_LIMIT]

    return {
        "name": _first_text(source, "name", "projectName", "title"),
        "client": _first_text(source, "client", "clientName", "product"),
        "description": description,
        "technologies": _first_text(source, "technologies", "technologyStack", "techStack", "tools", "environment"),
        "developmentTools": _first_text(source, "developmentTools", "development_tools", "devTools"),
        "teamSize": _first_text(source, "teamSize", "team_size"),
        "responsibilities": responsibilities,
        "link": _first_text(source, "link", "url"),
    }


def _normalize_uploaded_experience(exp: Any) -> Dict[str, Any]:
    source = exp if isinstance(exp, dict) else {}
    return {
        "company": _first_text(source, "company", "employer", "organization"),
        "title": _first_text(source, "title", "position", "role", "designation"),
        "location": _first_text(source, "location", "city"),
        "startDate": _first_text(source, "startDate", "start_date", "from"),
        "endDate": _first_text(source, "endDate", "end_date", "to"),
        "responsibilities": _first_text_list(
            source,
            "responsibilities",
            "responsibility",
            "bullets",
            "bulletPoints",
            "achievements",
            "highlights",
            "contributions",
        )[:PROFESSIONAL_EXPERIENCE_BULLET_LIMIT],
    }


def _normalize_uploaded_resume_data(data: Dict[str, Any]) -> Dict[str, Any]:
    data = data if isinstance(data, dict) else {}
    header = data.get("header") if isinstance(data.get("header"), dict) else {}
    expertise = data.get("expertise") if isinstance(data.get("expertise"), dict) else {}
    skills = data.get("skills") if isinstance(data.get("skills"), dict) else {}

    return {
        **data,
        "header": {
            "fullName": _first_text(header, "fullName", "name"),
            "designation": _first_text(header, "designation", "title"),
            "email": _first_text(header, "email"),
            "phone": _first_text(header, "phone"),
            "location": _first_text(header, "location"),
            "linkedin": _first_text(header, "linkedin", "linkedIn"),
            "github": _first_text(header, "github", "gitHub"),
            "portfolio": _first_text(header, "portfolio", "website"),
        },
        "expertise": {
            "summary": _first_text(expertise, "summary", "profile", "objective"),
            "bulletPoints": _first_text_list(expertise, "bulletPoints", "bullets", "highlights"),
        },
        "skills": {"skills": _first_text(skills, "skills", "technicalSkills") or _clean_text(data.get("skills"))},
        "experiences": [
            _normalize_uploaded_experience(exp)
            for exp in _item_list(data.get("experiences") or data.get("experience"))
        ],
        "projects": [
            _normalize_uploaded_project(project)
            for project in _item_list(data.get("projects") or data.get("project"))
        ],
        "education": _item_list(data.get("education")),
        "awards": _item_list(data.get("awards")),
    }


def _extract_text_from_pdf(file_content: bytes) -> str:
    """Extract text from PDF file."""
    try:
        import PyPDF2
        from io import BytesIO
        
        pdf_file = BytesIO(file_content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        text = ""
        for page in pdf_reader.pages:
            text += page.extract_text() + "\n"
        return text
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract text from PDF: {str(e)}"
        )


def _extract_text_from_docx(file_content: bytes) -> str:
    """Extract text from DOCX file."""
    try:
        from docx import Document
        from io import BytesIO
        
        docx_file = BytesIO(file_content)
        doc = Document(docx_file)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        return text
    except Exception as e:
        raise HTTPException(
            status_code=400,
            detail=f"Failed to extract text from DOCX: {str(e)}"
        )


@router.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)) -> Dict[str, Any]:
    """
    Upload a resume file (PDF or DOCX) and extract structured data using AI.
    
    This endpoint accepts PDF or DOCX files, extracts text, and uses AI to parse
    the resume content into structured data matching the ResumeData schema.
    """
    # Validate file type
    if not file.filename:
        raise HTTPException(status_code=400, detail="No file provided")

    upload_filename = file.filename
    raw = ""
    parsed_ai_data: Optional[Dict[str, Any]] = None
    
    file_extension = upload_filename.lower().split(".")[-1]
    if file_extension not in ["pdf", "docx"]:
        raise HTTPException(
            status_code=400,
            detail="Invalid file type. Please upload a PDF or DOCX file."
        )
    
    # Check if OpenAI API key is configured
    api_key = os.getenv("OPENAI_API_KEY")
    if not api_key:
        raise HTTPException(
            status_code=503,
            detail="AI parse is not configured. Set OPENAI_API_KEY in backend/.env",
        )
    
    try:
        # Read file content
        file_content = await file.read()
        
        # Extract text based on file type
        if file_extension == "pdf":
            extracted_text = _extract_text_from_pdf(file_content)
        else:  # docx
            extracted_text = _extract_text_from_docx(file_content)
        
        if not extracted_text.strip():
            raise HTTPException(
                status_code=400,
                detail="Could not extract text from the uploaded file. Please ensure the file contains readable text."
            )
        
        # Use AI to parse the extracted text
        from openai import OpenAI
        
        base_url = os.getenv("OPENAI_API_BASE", "https://openrouter.ai/api/v1")
        client = OpenAI(api_key=api_key, base_url=base_url)
        
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": RESUME_SCHEMA_PROMPT},
                {"role": "user", "content": extracted_text.strip()[:12000]},
            ],
            temperature=0.2,
        )
        
        raw = resp.choices[0].message.content or "{}"
        data = _extract_json_from_response(raw)
        parsed_ai_data = json.loads(json.dumps(data))
        
    except json.JSONDecodeError as e:
        _write_ai_output_log(
            upload_filename=upload_filename,
            raw_ai_response=raw,
            error=f"AI returned invalid JSON: {e}",
        )
        raise HTTPException(status_code=502, detail=f"AI returned invalid JSON: {e}")
    except HTTPException:
        raise
    except Exception as e:
        raise HTTPException(status_code=502, detail=f"Failed to process resume: {str(e)}")
    
    data = _normalize_uploaded_resume_data(data)

    _write_ai_output_log(
        upload_filename=upload_filename,
        raw_ai_response=raw,
        parsed_data=parsed_ai_data,
        returned_data=data,
    )
    
    return data
