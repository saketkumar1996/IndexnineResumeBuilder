"""
Integration tests for resume rendering (HTML and DOCX generation)
Tests template rendering, output consistency, and document generation
"""

import pytest
from models.resume_models import ResumeModel
from rendering.resume_renderer import ResumeRenderer
import tempfile
import os
from docx import Document


class TestResumeRendering:
    """Integration tests for resume rendering functionality"""
    
    @pytest.fixture
    def valid_resume_model(self):
        """Complete valid resume model for testing"""
        return ResumeModel(**{
            "header": {
                "name": "Jane Smith",
                "title": "Senior Full Stack Developer",
                "email": "jane.smith@example.com",
                "phone": "+1 (555) 987-6543",
                "location": "Seattle, WA"
            },
            "expertise": {
                "summary": "Highly skilled full stack developer with over 8 years of experience in designing and implementing scalable web applications. Expertise in modern JavaScript frameworks, cloud technologies, and microservices architecture. Proven track record of leading development teams and delivering high-quality software solutions. Strong background in agile methodologies, test-driven development, and DevOps practices. Passionate about clean code, performance optimization, and user experience. Experienced in working with cross-functional teams to translate business requirements into technical solutions. Committed to continuous learning and staying current with emerging technologies and industry best practices."
            },
            "skills": {
                "skills": "JavaScript, TypeScript, React, Node.js, Python, Django, AWS, Docker, Kubernetes, PostgreSQL, MongoDB, Redis, GraphQL, REST APIs, Git, Jenkins, Terraform"
            },
            "experience": [
                {
                    "company": "TechCorp Solutions",
                    "position": "Senior Full Stack Developer",
                    "start_date": "MAR 2020",
                    "end_date": "Present",
                    "responsibilities": [
                        "Led development of microservices architecture serving 500K+ daily active users",
                        "Mentored team of 5 junior developers and established code review processes",
                        "Implemented CI/CD pipelines reducing deployment time by 60%",
                        "Architected and built real-time notification system using WebSockets and Redis"
                    ]
                },
                {
                    "company": "StartupXYZ",
                    "position": "Full Stack Developer",
                    "start_date": "JUN 2018",
                    "end_date": "FEB 2020",
                    "responsibilities": [
                        "Developed responsive web applications using React and Node.js",
                        "Built RESTful APIs and GraphQL endpoints for mobile and web clients",
                        "Optimized database queries resulting in 40% performance improvement"
                    ]
                }
            ],
            "projects": [
                {
                    "name": "E-commerce Platform Redesign",
                    "description": "Led complete redesign of e-commerce platform handling $2M+ monthly transactions",
                    "technologies": "React, Next.js, Node.js, PostgreSQL, Redis, AWS",
                    "start_date": "JAN 2021",
                    "end_date": "JUN 2021"
                },
                {
                    "name": "Real-time Analytics Dashboard",
                    "description": "Built comprehensive analytics dashboard providing real-time insights",
                    "technologies": "Vue.js, Python, FastAPI, InfluxDB, Docker",
                    "start_date": "SEP 2020",
                    "end_date": "DEC 2020"
                }
            ],
            "education": [
                {
                    "institution": "University of Washington",
                    "degree": "Master of Science",
                    "field_of_study": "Computer Science",
                    "graduation_date": "JUN 2018",
                    "gpa": "3.9"
                },
                {
                    "institution": "Seattle University",
                    "degree": "Bachelor of Science",
                    "field_of_study": "Software Engineering",
                    "graduation_date": "MAY 2016",
                    "gpa": "3.7"
                }
            ],
            "awards": [
                {
                    "title": "Employee of the Year",
                    "organization": "TechCorp Solutions",
                    "date": "DEC 2022",
                    "description": "Recognized for outstanding technical leadership"
                }
            ]
        })
    
    @pytest.fixture
    def renderer(self):
        """Resume renderer instance"""
        return ResumeRenderer(template_path="templates")
    
    def test_html_rendering_contains_all_sections(self, renderer, valid_resume_model):
        """Test that HTML rendering includes all resume sections"""
        html_content = renderer.render_html(valid_resume_model)
        
        # Check header information
        assert "Jane Smith" in html_content
        assert "Senior Full Stack Developer" in html_content
        assert "jane.smith@example.com" in html_content
        assert "+1 (555) 987-6543" in html_content
        assert "Seattle, WA" in html_content
        
        # Check expertise section
        assert "Highly skilled full stack developer" in html_content
        
        # Check skills section
        assert "JavaScript, TypeScript, React" in html_content
        
        # Check experience section
        assert "TechCorp Solutions" in html_content
        assert "StartupXYZ" in html_content
        assert "Led development of microservices" in html_content
        
        # Check projects section
        assert "E-commerce Platform Redesign" in html_content
        assert "Real-time Analytics Dashboard" in html_content
        
        # Check education section
        assert "University of Washington" in html_content
        assert "Seattle University" in html_content
        
        # Check awards section
        assert "Employee of the Year" in html_content
    
    def test_html_rendering_proper_structure(self, renderer, valid_resume_model):
        """Test that HTML has proper structure and formatting"""
        html_content = renderer.render_html(valid_resume_model)
        
        # Should contain proper HTML structure
        assert "<html" in html_content
        assert "<head>" in html_content
        assert "<body>" in html_content
        assert "</html>" in html_content
        
        # Should contain CSS styling
        assert "style" in html_content.lower()
        
        # Should have proper section headers
        assert any(header in html_content.lower() for header in ["expertise", "summary"])
        assert "skills" in html_content.lower()
        assert "experience" in html_content.lower()
        assert "projects" in html_content.lower()
        assert "education" in html_content.lower()
        assert "awards" in html_content.lower()
    
    def test_docx_generation_creates_valid_document(self, renderer, valid_resume_model):
        """Test that DOCX generation creates a valid Word document"""
        doc = renderer.generate_docx(valid_resume_model)
        
        # Should be a valid Document object
        assert isinstance(doc, Document)
        
        # Save to temporary file and verify it can be opened
        with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
            doc.save(tmp_file.name)
            
            # Verify file exists and has content
            assert os.path.exists(tmp_file.name)
            assert os.path.getsize(tmp_file.name) > 0
            
            # Verify it can be opened as a Word document
            reopened_doc = Document(tmp_file.name)
            assert len(reopened_doc.paragraphs) > 0
            
            # Clean up
            os.unlink(tmp_file.name)
    
    def test_docx_contains_all_resume_content(self, renderer, valid_resume_model):
        """Test that DOCX document contains all resume content"""
        doc = renderer.generate_docx(valid_resume_model)
        
        # Extract all text from document
        full_text = []
        for paragraph in doc.paragraphs:
            full_text.append(paragraph.text)
        
        # Also check tables (if any)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    full_text.append(cell.text)
        
        document_text = " ".join(full_text)
        
        # Check that all key content is present
        assert "Jane Smith" in document_text
        assert "Senior Full Stack Developer" in document_text
        assert "jane.smith@example.com" in document_text
        assert "TechCorp Solutions" in document_text
        assert "University of Washington" in document_text
        assert "E-commerce Platform Redesign" in document_text
        assert "Employee of the Year" in document_text
    
    def test_html_docx_content_consistency(self, renderer, valid_resume_model):
        """Test that HTML and DOCX contain consistent content"""
        html_content = renderer.render_html(valid_resume_model)
        doc = renderer.generate_docx(valid_resume_model)
        
        # Extract text from DOCX
        docx_text = []
        for paragraph in doc.paragraphs:
            docx_text.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    docx_text.append(cell.text)
        
        docx_content = " ".join(docx_text)
        
        # Key content should be present in both formats
        key_content = [
            "Jane Smith",
            "Senior Full Stack Developer",
            "TechCorp Solutions",
            "University of Washington",
            "E-commerce Platform Redesign",
            "Employee of the Year"
        ]
        
        for content in key_content:
            assert content in html_content, f"'{content}' missing from HTML"
            assert content in docx_content, f"'{content}' missing from DOCX"
    
    def test_partial_html_rendering_with_errors(self, renderer):
        """Test partial HTML rendering when validation errors exist"""
        partial_data = {
            "header": {
                "name": "John Doe",
                "title": "Software Engineer",
                "email": "john@example.com",
                "phone": "+1 555-123-4567",
                "location": "San Francisco, CA"
            },
            "expertise": {
                "summary": "Too short"  # Invalid
            }
        }
        
        partial_html = renderer.render_partial_html(partial_data)
        
        # Should contain available data
        assert "John Doe" in partial_html
        assert "Software Engineer" in partial_html
        
        # Should indicate incomplete/error state
        assert any(indicator in partial_html.lower() for indicator in ["incomplete", "error", "validation"])
    
    def test_empty_optional_sections_handling(self, renderer):
        """Test rendering with empty optional sections"""
        minimal_resume = ResumeModel(**{
            "header": {
                "name": "John Doe",
                "title": "Software Engineer",
                "email": "john@example.com",
                "phone": "+1 555-123-4567",
                "location": "San Francisco, CA"
            },
            "expertise": {
                "summary": "Experienced software engineer with expertise in web development and cloud technologies. Skilled in multiple programming languages and frameworks. Strong problem-solving abilities and excellent communication skills. Passionate about creating efficient and scalable solutions. Committed to continuous learning and professional development. Proven track record of successful project delivery. Collaborative team player with leadership experience. Dedicated to writing clean and maintainable code. Experienced in agile development methodologies and best practices. Always eager to take on new challenges and learn emerging technologies. Focused on delivering high-quality software solutions."
            },
            "skills": {
                "skills": "Python, JavaScript, React, Node.js"
            },
            "experience": [{
                "company": "Tech Corp",
                "position": "Developer",
                "start_date": "JAN 2020",
                "end_date": "Present",
                "responsibilities": [
                    "Developed web applications",
                    "Collaborated with team members",
                    "Implemented best practices"
                ]
            }],
            "projects": [{
                "name": "Web App",
                "description": "Built a web application",
                "technologies": "Python, React",
                "start_date": "JAN 2020",
                "end_date": "DEC 2020"
            }],
            "education": [{
                "institution": "University",
                "degree": "Bachelor",
                "field_of_study": "Computer Science",
                "graduation_date": "MAY 2020"
            }],
            "awards": []  # Empty optional section
        })
        
        html_content = renderer.render_html(minimal_resume)
        doc = renderer.generate_docx(minimal_resume)
        
        # Should render successfully without awards section
        assert "John Doe" in html_content
        assert len(doc.paragraphs) > 0
        
        # Awards section should either be omitted or show as empty
        # (implementation dependent)
    
    def test_special_characters_handling(self, renderer):
        """Test rendering with special characters and international names"""
        special_char_resume = ResumeModel(**{
            "header": {
                "name": "José María O'Connor-Smith",
                "title": "Software Engineer & Data Scientist",
                "email": "jose.maria@example.com",
                "phone": "+1 (555) 123-4567",
                "location": "São Paulo, Brazil"
            },
            "expertise": {
                "summary": "Experienced software engineer with expertise in web development and data science. Skilled in multiple programming languages including Python, R, and JavaScript. Strong background in machine learning and statistical analysis. Passionate about creating innovative solutions that drive business value. Committed to continuous learning and professional development. Proven track record of successful project delivery in international environments. Collaborative team player with excellent communication skills. Dedicated to writing clean, maintainable code. Experienced in agile methodologies and cross-cultural team collaboration. Always eager to tackle complex challenges."
            },
            "skills": {
                "skills": "Python, R, JavaScript, SQL, Machine Learning, Data Analysis"
            },
            "experience": [{
                "company": "Tech Corp & Associates",
                "position": "Senior Developer/Analyst",
                "start_date": "JAN 2020",
                "end_date": "Present",
                "responsibilities": [
                    "Developed ML models with 95% accuracy",
                    "Led cross-functional teams (5-10 members)",
                    "Implemented CI/CD pipelines & automated testing"
                ]
            }],
            "projects": [{
                "name": "AI-Powered Analytics Platform",
                "description": "Built ML platform processing 1M+ records/day",
                "technologies": "Python, TensorFlow, AWS, Docker",
                "start_date": "JAN 2021",
                "end_date": "DEC 2021"
            }],
            "education": [{
                "institution": "Universidade de São Paulo",
                "degree": "Master's in Computer Science",
                "field_of_study": "Artificial Intelligence & Machine Learning",
                "graduation_date": "DEC 2019"
            }],
            "awards": []
        })
        
        html_content = renderer.render_html(special_char_resume)
        doc = renderer.generate_docx(special_char_resume)
        
        # Should handle special characters correctly
        assert "José María O'Connor-Smith" in html_content
        assert "São Paulo, Brazil" in html_content
        assert "Tech Corp & Associates" in html_content
        assert "Universidade de São Paulo" in html_content
        
        # DOCX should also handle special characters
        docx_text = " ".join([p.text for p in doc.paragraphs])
        assert "José María" in docx_text
        assert "São Paulo" in docx_text
    
    def test_rendering_performance(self, renderer, valid_resume_model):
        """Test that rendering completes within reasonable time"""
        import time
        
        # HTML rendering should be fast
        start_time = time.time()
        html_content = renderer.render_html(valid_resume_model)
        html_time = time.time() - start_time
        
        assert html_time < 1.0  # Should complete within 1 second
        assert len(html_content) > 1000  # Should generate substantial content
        
        # DOCX generation should also be reasonably fast
        start_time = time.time()
        doc = renderer.generate_docx(valid_resume_model)
        docx_time = time.time() - start_time
        
        assert docx_time < 5.0  # Should complete within 5 seconds
        assert len(doc.paragraphs) > 10  # Should generate multiple paragraphs