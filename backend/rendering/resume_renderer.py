"""
Resume rendering module for HTML and DOCX generation
Handles template rendering and document creation with consistent output
"""

from jinja2 import Environment, FileSystemLoader
from docx import Document
from docx.shared import Inches
from models.resume_models import ResumeModel
from typing import Dict, Any
import os


class ResumeRenderer:
    """Handles HTML and DOCX rendering for resume data"""
    
    def __init__(self, template_path: str = "templates"):
        """Initialize renderer with template path"""
        self.template_path = template_path
        self.jinja_env = Environment(loader=FileSystemLoader(template_path))
    
    def render_html(self, resume_data: ResumeModel) -> str:
        """Render resume data to HTML using Jinja2 template"""
        try:
            template = self.jinja_env.get_template("resume.html")
            return template.render(resume=resume_data)
        except Exception as e:
            # Fallback to basic HTML structure
            return self._generate_basic_html(resume_data)
    
    def render_partial_html(self, partial_data: Dict[str, Any]) -> str:
        """Render partial resume data with placeholders for missing sections"""
        try:
            template = self.jinja_env.get_template("resume_partial.html")
            return template.render(data=partial_data)
        except Exception:
            # Fallback to basic partial HTML
            return self._generate_partial_html(partial_data)
    
    def generate_docx(self, resume_data: ResumeModel) -> Document:
        """Generate DOCX document from resume data"""
        doc = Document()
        
        # Header section
        header_para = doc.add_paragraph()
        header_para.alignment = 1  # Center alignment
        name_run = header_para.add_run(resume_data.header.name)
        name_run.bold = True
        name_run.font.size = 14
        
        header_para.add_run(f"\n{resume_data.header.title}")
        header_para.add_run(f"\n{resume_data.header.email} | {resume_data.header.phone}")
        header_para.add_run(f"\n{resume_data.header.location}")
        
        # Expertise section
        doc.add_heading("EXPERTISE", level=2)
        doc.add_paragraph(resume_data.expertise.summary)
        
        # Skills section
        doc.add_heading("SKILLS", level=2)
        doc.add_paragraph(resume_data.skills.skills)
        
        # Experience section
        doc.add_heading("EXPERIENCE", level=2)
        for exp in resume_data.experience:
            exp_para = doc.add_paragraph()
            exp_para.add_run(f"{exp.company} - {exp.position}").bold = True
            exp_para.add_run(f"\n{exp.start_date} - {exp.end_date or 'Present'}")
            
            for responsibility in exp.responsibilities:
                doc.add_paragraph(f"• {responsibility}")
        
        # Projects section
        doc.add_heading("PROJECT EXPERIENCE", level=2)
        for project in resume_data.projects:
            proj_para = doc.add_paragraph()
            proj_para.add_run(project.name).bold = True
            proj_para.add_run(f"\n{project.start_date} - {project.end_date or 'Present'}")
            doc.add_paragraph(project.description)
            doc.add_paragraph(f"Technologies: {project.technologies}")
        
        # Education section
        doc.add_heading("EDUCATION", level=2)
        for edu in resume_data.education:
            edu_para = doc.add_paragraph()
            edu_para.add_run(f"{edu.institution} - {edu.degree}").bold = True
            edu_para.add_run(f"\n{edu.field_of_study}")
            edu_para.add_run(f"\nGraduated: {edu.graduation_date}")
            if edu.gpa:
                edu_para.add_run(f" | GPA: {edu.gpa}")
        
        # Awards section (if any)
        if resume_data.awards:
            doc.add_heading("AWARDS", level=2)
            for award in resume_data.awards:
                award_para = doc.add_paragraph()
                award_para.add_run(f"{award.title} - {award.organization}").bold = True
                award_para.add_run(f"\n{award.date}")
                if award.description:
                    award_para.add_run(f"\n{award.description}")
        
        return doc
    
    def _generate_basic_html(self, resume_data: ResumeModel) -> str:
        """Generate basic HTML structure as fallback"""
        html = f"""
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <meta name="viewport" content="width=device-width, initial-scale=1.0">
            <title>Resume Preview</title>
            <style>
                body {{ font-family: 'Times New Roman', serif; margin: 0.5in; }}
                .header {{ text-align: center; margin-bottom: 20px; }}
                .name {{ font-size: 18px; font-weight: bold; }}
                .section-header {{ font-weight: bold; margin-top: 20px; margin-bottom: 10px; }}
            </style>
        </head>
        <body>
            <div class="header">
                <div class="name">{resume_data.header.name}</div>
                <div>{resume_data.header.title}</div>
                <div>{resume_data.header.email} | {resume_data.header.phone}</div>
                <div>{resume_data.header.location}</div>
            </div>
            
            <div class="section-header">EXPERTISE</div>
            <p>{resume_data.expertise.summary}</p>
            
            <div class="section-header">SKILLS</div>
            <p>{resume_data.skills.skills}</p>
            
            <div class="section-header">EXPERIENCE</div>
        """
        
        for exp in resume_data.experience:
            html += f"""
            <div>
                <strong>{exp.company} - {exp.position}</strong><br>
                {exp.start_date} - {exp.end_date or 'Present'}<br>
                <ul>
            """
            for resp in exp.responsibilities:
                html += f"<li>{resp}</li>"
            html += "</ul></div>"
        
        html += """
            </body>
            </html>
        """
        return html
    
    def _generate_partial_html(self, partial_data: Dict[str, Any]) -> str:
        """Generate partial HTML with available data"""
        html = """
        <!DOCTYPE html>
        <html lang="en">
        <head>
            <meta charset="UTF-8">
            <title>Resume Preview - Incomplete</title>
            <style>
                body { font-family: 'Times New Roman', serif; margin: 0.5in; }
                .placeholder { color: #888; font-style: italic; }
                .validation-error { color: red; margin: 10px 0; }
            </style>
        </head>
        <body>
            <div class="validation-error">Resume validation incomplete. Please fix errors and complete all required sections.</div>
        """
        
        if "header" in partial_data:
            header = partial_data["header"]
            html += f"""
            <div style="text-align: center; margin-bottom: 20px;">
                <div style="font-size: 18px; font-weight: bold;">{header.get('name', 'Name not provided')}</div>
                <div>{header.get('title', 'Title not provided')}</div>
            </div>
            """
        
        html += "</body></html>"
        return html